import PyPDF2
import docx
import io
import re
from typing import Optional

class PDFProcessor:
    def __init__(self):
        pass
    
    async def extract_text(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from PDF or DOCX files
        """
        try:
            if filename.lower().endswith('.pdf'):
                return await self._extract_from_pdf(file_content)
            elif filename.lower().endswith(('.docx', '.doc')):
                return await self._extract_from_docx(file_content)
            else:
                raise ValueError("Unsupported file format")
        except Exception as e:
            raise Exception(f"Failed to extract text: {str(e)}")
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file
        """
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        """
        # Remove extra whitespace and newlines
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\-@.,():/]', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def extract_sections(self, text: str) -> dict:
        """
        Extract different sections from resume text
        """
        sections = {
            'contact_info': '',
            'summary': '',
            'experience': '',
            'education': '',
            'skills': '',
            'projects': '',
            'certifications': ''
        }
        
        # Define section patterns
        section_patterns = {
            'contact_info': r'(?:contact|personal information|details).*?(?=\n[A-Z])',
            'summary': r'(?:summary|objective|profile).*?(?=\n[A-Z])',
            'experience': r'(?:experience|work history|employment).*?(?=\n[A-Z])',
            'education': r'(?:education|academic).*?(?=\n[A-Z])',
            'skills': r'(?:skills|technical skills|competencies).*?(?=\n[A-Z])',
            'projects': r'(?:projects|portfolio).*?(?=\n[A-Z])',
            'certifications': r'(?:certifications|certificates|licenses).*?(?=\n[A-Z])'
        }
        
        text_lower = text.lower()
        
        for section, pattern in section_patterns.items():
            match = re.search(pattern, text_lower, re.DOTALL | re.IGNORECASE)
            if match:
                sections[section] = match.group(0)
        
        return sections
    
    def extract_contact_info(self, text: str) -> dict:
        """
        Extract contact information from resume text
        """
        contact_info = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None
        }
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group(0)
        
        # Phone pattern
        phone_pattern = r'(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info['phone'] = phone_match.group(0)
        
        # LinkedIn pattern
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/pub/)([a-zA-Z0-9-]+)'
        linkedin_match = re.search(linkedin_pattern, text.lower())
        if linkedin_match:
            contact_info['linkedin'] = f"linkedin.com/in/{linkedin_match.group(1)}"
        
        # GitHub pattern
        github_pattern = r'(?:github\.com/)([a-zA-Z0-9-]+)'
        github_match = re.search(github_pattern, text.lower())
        if github_match:
            contact_info['github'] = f"github.com/{github_match.group(1)}"
        
        return contact_info